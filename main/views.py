import docx
import pandas as pd
import numpy as np
from django.shortcuts import render
from docx.shared import Inches
from matplotlib import pyplot as plt
from scipy import stats
from sklearn.metrics import mean_absolute_error

def index(request):
    if request.method == 'POST':
        file = request.FILES['file']

        alpha = float(request.POST.get('value'))

        if file.name.endswith('.csv'):
            df = pd.read_csv(file)
            df = df.dropna()

            df['X'] = pd.to_numeric(df['X'], errors='coerce')
            df['Y'] = pd.to_numeric(df['Y'], errors='coerce')
            df = df.dropna()
            df = df.round(2)

            table = df.to_html(
                classes='table table-striped table-hover table-bordered table-sm table-responsive text-center')

            # вычисление среднеквадратичного отклонения по Y и X
            std_y = np.std(df['Y'])
            std_x = np.std(df['X'])

            # вычисление линейного коэффициента корреляции
            corr_coef = df['X'].corr(df['Y'])

            # определение уровня корреляции между признаками x и y
            corr_level = 'слабая'
            if abs(corr_coef) > 0.7:
                corr_level = 'сильная'
            elif abs(corr_coef) > 0.3:
                corr_level = 'средняя'

            # вычисление средней ошибки коэффициента корреляции
            stderr_corr_coef = (1 - corr_coef ** 2) / (np.sqrt(len(df)))

            # нужен критерий Стьюдента
            df_len = len(df)
            t_alpha = stats.t.ppf(1 - alpha / 2, df_len - 2)

            # проверка коэффициента корреляции на значимость
            t_score = abs(corr_coef) * np.sqrt(len(df) - 2) / np.sqrt(1 - corr_coef ** 2)
            p_value = 2 * (1 - stats.t.cdf(t_score, df=df.shape[0] - 2))

            if len(df)>50:
                p_value = (1-corr_coef**2)/(np.sqrt(len(df)))
            elif len(df)<30:
                p_value = np.sqrt(1-corr_coef**2)/(np.sqrt(len(df)-2))

            significance = p_value < alpha
            if significance:
                significance = 'Значим'
            else:
                significance ='Не значим'

            # Коэффициент Спирмена
            spearman_corr = df['X'].corr(df['Y'], method='spearman')

            # Уровень между признаками
            corr_level2 = 'слабая'
            if abs(spearman_corr) > 0.7:
                corr_level2 = 'сильная'
            elif abs(spearman_corr) > 0.3:
                corr_level2 = 'средняя'

            # Эластичность

            mean_y = df['Y'].mean()
            mean_x = df['X'].mean()
            a1, a0 = np.polyfit(df['X'], df['Y'], 1)
            elasticity = a1 * (mean_x / mean_y)


            #Средняя ошибка аппроксимации
            y_pred = a0 + a1 * df['X']
            approx_error = mean_absolute_error(df['Y'], y_pred)

            # Общая дисперсия

            total_var = np.sum((df['Y'] - mean_y)**2)/df_len

            # Факторная дисперсия
            factor_var =np.sum((y_pred - mean_y)**2)/df_len

            #Остаточная дисперсия

            res_var = np.sum((df['Y'] - y_pred)**2)/df_len

            # Проверка общей дисперсии
            # Если общая дисперсия равна сумме факторной и остаточной дисперсии, проверка пройдена
            var_check = np.isclose(total_var, factor_var + res_var)
            if var_check:
                var_check = 'Верно'
            else:
                var_check = 'Не верно'

            # Теоретический коэффициент детерминации

            determination_coef = np.sqrt(factor_var**2/total_var**2)

            #Теоретическое корреляционное отношение
            corr_ratio = np.sqrt(determination_coef)

            #Зависимость между коррелированными отношениями
            rel_level = 'слабая'
            if abs(corr_ratio) > 0.7:
                rel_level = 'сильная'
            elif abs(corr_ratio) > 0.3:
                rel_level = 'средняя'

            #Средние ошибки параметров
            avg_error_a0 = np.sqrt(res_var)/np.sqrt(df_len-2)
            avg_error_a1 = np.sqrt(res_var) / (std_x * np.sqrt(df_len - 2))
            #Отношение коэффициента к его средней ошибке
            t_a0 = a0/avg_error_a0
            t_a1 = a1 / avg_error_a1

            if t_a0 > 3:
                check_t_a0 = 'Значим'
            else:
                check_t_a0 = 'Не значим'

            if t_a1 > 3:
                check_t_a1 = 'Значим'
            else:
                check_t_a1 = 'Не значим'

            #F-критерий Фишера
            f_score = determination_coef / (1 - determination_coef) * (len(df) - 2)

            #При заданном уровне значимости считает F табличное
            f_critical = stats.f.ppf(1 - alpha, 1, len(df) - 2)

            if f_score > f_critical:
                f_check = 'Значим'
            else:
                f_check = 'Не значим'

            # Убедимся, что X не равен 0 для гиперболической регрессии
            df = df[df['X'] != 0]

            # Построение уравнений регрессии
            x = df['X'].values
            y = df['Y'].values

            # Линейная регрессия
            linear_coefs = np.polyfit(x, y, 1)
            y_linear = np.polyval(linear_coefs, x)

            # Параболическая регрессия
            x_quadratic = np.linspace(x.min(), x.max(), 500)
            quadratic_coefs = np.polyfit(x, y, 2)
            y_quadratic = np.polyval(quadratic_coefs, x_quadratic)

            # Гиперболическая регрессия
            x_hyperbolic = np.linspace(x.min(), x.max(), 500)
            hyperbolic_coefs = np.polyfit(1 / x, y, 1)
            y_hyperbolic = np.polyval(hyperbolic_coefs, 1 / x_hyperbolic)
            plt.rcParams['lines.markersize'] = 0.8
            # Построение графиков
            fig, ax = plt.subplots(3, 1, figsize=(10, 15))

            ax[0].scatter(x, y, color='blue')
            ax[0].plot(x, y_linear, color='red')
            ax[0].set_title('Линейная регрессия')
            ax[0].set_xlabel('Курс $')
            ax[0].set_ylabel('Объем млдр.м3')

            ax[1].scatter(x, y, color='blue')
            ax[1].plot(x_quadratic, y_quadratic, color='red')
            ax[1].set_title('Параболическая регрессия')
            ax[1].set_xlabel('Курс $')
            ax[1].set_ylabel('Объем млдр.м3')

            ax[2].scatter(x, y, color='blue')
            ax[2].plot(x_hyperbolic, y_hyperbolic, color='red')
            ax[2].set_title('Гиперболическая регрессия')
            ax[2].set_xlabel('Курс $')
            ax[2].set_ylabel('Объем млдр.м3')

            plt.tight_layout()
            plt.savefig('plot.svg')
            plt.savefig('plot.png')

            #Сохранение в текстовый формат
            doc = docx.Document()  # create
            table_docx = doc.add_table(rows=len(df) + 1,
                                  cols=2)
            table_docx.cell(0, 0).text = 'Курс'
            table_docx.cell(0, 1).text = 'Газ млрд.кб. м3'
            for i in range(len(df)):
                table_docx.cell(i + 1, 0).text = str(df.iloc[i, 0])
                table_docx.cell(i + 1, 1).text = str(df.iloc[i, 1])
            doc.add_page_break()
            textfile = ('Уровень значимости: '+ str(alpha)+"\n"+'Критерий Стьюдента: '+str(t_alpha)+"\n"+'Коэффициент Спирмена: '+str(spearman_corr)+"\n"
                              'Зависимость между коррелированными отношениями: '+str(corr_level)+"\n"+'Уровень между признаками: '+str(corr_level2)+"\n"
                              'Эластичность: '+str(elasticity)+"\n"
                              'Средняя ошибка аппроксимации: '+str(approx_error)+"\n"
                              'Общая дисперсия: '+str(total_var)+"\n"
                              'Факторная дисперсия: '+str(factor_var)+"\n"
                              'Остаточная дисперсия: '+str(res_var)+"\n"
                              'Проверка общей дисперсии: '+str(var_check)+"\n"
                              'Теоретический коэффициент детерминации: '+str(determination_coef)+"\n"
                              'Теоретическое корреляционное отношение: '+str(corr_ratio)+"\n"
                              'Среднеквадратичное отклонение Y:'+str(std_y)+"\n"
                              'Среднеквадратичное отклонение X:'+str(std_x)+"\n"
                              'Линейный коэффициент корреляции'+str(corr_coef)+"\n"
                              'Средней ошибки коэффициента корреляции: '+str(stderr_corr_coef)+"\n"
                              'Коэффициента корреляции на значимость: '+str(significance)+"\n"
                              'Зависимость между коррелированными отношениями: '+str(rel_level)+"\n"
                              'F - критерий Фишера: '+str(f_score)+"\n"
                              'F - табличное: '+str(f_critical) +"\n" +str(f_check)+"\n"
                              'Средние ошибки параметров:'+'a0: '+str(avg_error_a0)+'a1: '+str(avg_error_a1)+ "\n"
                              't a0'+str(t_a0) +
                              ' '+str(check_t_a0) + "\n"
                              't a1'+str(t_a1) +
                              ' '+str(check_t_a1) + "\n")
            doc.add_paragraph(textfile
                              )
            doc.add_page_break()
            doc.add_picture('plot.png', width=Inches(5.83), height=Inches(8.27))

            doc.save('document.docx')

            return render(request, 'success.html', {
                'table': table,
                'alpha': alpha,
                't_alpha': t_alpha,
                'spearman_corr': spearman_corr,
                'corr_level': corr_level,
                'corr_level2': corr_level2,
                'elasticity': elasticity,
                'approx_error': approx_error,
                'total_var': total_var,
                'factor_var': factor_var,
                'res_var': res_var,
                'var_check': var_check,
                'determination_coef': determination_coef,
                'corr_ratio': corr_ratio,
                'std_y': std_y,
                'std_x': std_x,
                'corr_coef': corr_coef,
                'stderr_corr_coef': stderr_corr_coef,
                't_score': t_score,
                'p_value': p_value,
                'significance': significance,
                'rel_level': rel_level,
                'f_score': f_score,
                'f_critical': f_critical,
                't_a0': t_a0,
                't_a1': t_a1,
                'avg_error_a1': avg_error_a1,
                'avg_error_a0': avg_error_a0,
                'check_t_a0': check_t_a0,
                'check_t_a1': check_t_a1,
                'f_check':f_check,
            })

        else:
            return render(request, 'upload.html')
    else:
        return render(request, 'upload.html')


